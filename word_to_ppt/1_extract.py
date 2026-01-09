#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
å¾ Word æ–‡ä»¶ä¸­æå–è—è‰²æ–‡å­—ä¸¦è½‰æ›ç‚º PPT æ ¼å¼
Extract blue text from Word document and convert to PPT format
"""

from docx import Document
from docx.shared import RGBColor
import sys
import os
import traceback
from datetime import datetime


class BlueTextExtractor:
    """ç‰¹å®šé¡è‰²æ–‡å­—æå–å™¨"""
    
    def __init__(self, target_color=None, tolerance=50):
        """
        åˆå§‹åŒ–æå–å™¨
        
        Args:
            target_color: ç›®æ¨™é¡è‰² (r, g, b) æˆ– "#RRGGBB"ï¼Œé è¨­ç‚ºè—è‰²
            tolerance: é¡è‰²å®¹å·®ï¼ˆ0-255ï¼‰
        """
        self.tolerance = tolerance
        self.extracted_text = []
        self.variables = {}  # å„²å­˜è‡ªå‹•æå–çš„è®Šæ•¸
        
        # è¨­å®šç›®æ¨™é¡è‰²ï¼ˆé è¨­è—è‰²ï¼‰
        if target_color is None:
            self.target_color = (0, 0, 255)  # é è¨­è—è‰²
        elif isinstance(target_color, str) and target_color.startswith('#'):
            # 16é€²ä½æ ¼å¼è½‰æ›
            hex_color = target_color.lstrip('#')
            self.target_color = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        elif isinstance(target_color, tuple) and len(target_color) == 3:
            self.target_color = target_color
        else:
            raise ValueError("target_color å¿…é ˆæ˜¯ (r, g, b) tuple æˆ– '#RRGGBB' æ ¼å¼")
    
    def is_target_color(self, rgb):
        """
        åˆ¤æ–·é¡è‰²æ˜¯å¦ç‚ºç›®æ¨™é¡è‰²ï¼ˆåœ¨å®¹å·®ç¯„åœå…§ï¼‰
        
        Args:
            rgb: RGBColor ç‰©ä»¶æˆ– tuple (r, g, b)
        
        Returns:
            bool: æ˜¯å¦ç‚ºç›®æ¨™é¡è‰²
        """
        if rgb is None:
            return False
        
        # ç²å– RGB å€¼
        if isinstance(rgb, RGBColor):
            r, g, b = rgb
        elif isinstance(rgb, tuple) and len(rgb) == 3:
            r, g, b = rgb
        else:
            return False
        
        # åˆ¤æ–·æ˜¯å¦åœ¨ç›®æ¨™é¡è‰²çš„å®¹å·®ç¯„åœå…§
        target_r, target_g, target_b = self.target_color
        return (abs(r - target_r) <= self.tolerance and
                abs(g - target_g) <= self.tolerance and
                abs(b - target_b) <= self.tolerance)
    
    # ä¿ç•™èˆŠæ–¹æ³•åç¨±ä»¥ç¶­æŒå‘ä¸‹ç›¸å®¹
    def is_blue(self, rgb):
        """å‘ä¸‹ç›¸å®¹çš„æ–¹æ³•ï¼Œå¯¦éš›èª¿ç”¨ is_target_color"""
        return self.is_target_color(rgb)
    
    def extract_from_paragraph(self, paragraph):
        """
        å¾æ®µè½ä¸­æå–è—è‰²æ–‡å­—
        
        Args:
            paragraph: docx æ®µè½ç‰©ä»¶
        
        Returns:
            str: æå–çš„è—è‰²æ–‡å­—ï¼ˆå¦‚æœæœ‰ï¼‰
        """
        blue_text = []
        
        for run in paragraph.runs:
            # æª¢æŸ¥æ–‡å­—é¡è‰²
            if run.font.color and run.font.color.type == 1:  # RGB é¡è‰²
                rgb = run.font.color.rgb
                if self.is_blue(rgb):
                    text = run.text.strip()
                    if text:
                        blue_text.append(text)
        
        return ' '.join(blue_text) if blue_text else None
    
    def extract_variables(self, docx_path):
        """è‡ªå‹•æå–æ–‡ä»¶è®Šæ•¸ï¼ˆæ—¥æœŸã€ç¦®æ‹œé¡å‹ã€ä¸»é¡Œã€ç¶“æ–‡ï¼‰"""
        import re
        
        doc = Document(docx_path)
        date = "2026å¹´1æœˆ1æ—¥"
        service_type = "é€±ä¸‰ç¦®æ‹œ"
        title = "æˆ‘æ˜¯ä¸»é¡Œ"
        verse_refs = "ã€ç®´è¨€27ç« 12ç¯€ã€è©©ç¯‡46ç¯‡1ç¯€ã€‘"
        verses = []
        
        all_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # 1. æå–æ—¥æœŸå’Œç¦®æ‹œé¡å‹
        for text in all_paragraphs:
            if 'å¹´' in text and 'æœˆ' in text and 'æ—¥' in text:
                date_match = re.search(r'(\d{4}å¹´\d{1,2}æœˆ\d{1,2}æ—¥)', text)
                if date_match:
                    date = date_match.group(1)
                
                if 'é€±' in text or 'ç¦®æ‹œ' in text:
                    for day in ['é€±ä¸€', 'é€±äºŒ', 'é€±ä¸‰', 'é€±å››', 'é€±äº”', 'é€±å…­', 'ä¸»æ—¥', 'é€±æ—¥']:
                        if day in text:
                            service_type = f"{day}ç¦®æ‹œ" if day != 'ä¸»æ—¥' else 'ä¸»æ—¥ç¦®æ‹œ'
                            break
                break
        
        # 2. æå–ä¸»é¡Œ
        title_lines = []
        for text in all_paragraphs:
            if 'ã€ˆ' in text or 'ã€' in text:
                break
            if 'å¹´' not in text and 'æœˆ' not in text:
                title_lines.append(text)
        
        if title_lines:
            title = ' '.join(title_lines[:3])
        
        # 3. æå–ç¶“æ–‡
        verse_list = []
        for text in all_paragraphs:
            if text.startswith('ã€ˆ') and 'ã€‰' in text:
                verse_ref = text.split('ã€‰')[0].lstrip('ã€ˆ')
                verse_list.append(verse_ref)
                
                if 'ã€‰' in text:
                    verse_content = text.split('ã€‰', 1)[1].strip()
                    verses.append(f"ã€ˆ{verse_ref}ã€‰{verse_content}")
        
        if verse_list:
            verse_refs = 'ã€' + 'ã€'.join(verse_list) + 'ã€‘'
        
        self.variables = {'æ—¥æœŸ': date, 'ç¦®æ‹œé¡å‹': service_type, 'ä¸»é¡Œ': title, 'ç¶“æ–‡ç« ç¯€': verse_refs}
        for i, verse in enumerate(verses, 1):
            self.variables[f'ç¶“æ–‡{i}'] = verse
        
        print(f"âœ… è‡ªå‹•æå–è®Šæ•¸:")
        print(f"  æ—¥æœŸ: {date}")
        print(f"  ç¦®æ‹œé¡å‹: {service_type}")
        print(f"  ä¸»é¡Œ: {title[:50]}...")
        print(f"  ç¶“æ–‡: {len(verses)} å€‹")
    
    def extract_from_docx(self, docx_path):
        """å¾ Word æ–‡ä»¶ä¸­æå–æ‰€æœ‰è—è‰²æ–‡å­—ï¼ˆé€£çºŒçš„è—è‰²æ®µè½æœƒåˆä½µï¼‰"""
        try:
            # å…ˆæå–è®Šæ•¸
            self.extract_variables(docx_path)
            
            doc = Document(docx_path)
            self.extracted_text = []
            current_group = []
            
            for paragraph in doc.paragraphs:
                blue_text = self.extract_from_paragraph(paragraph)
                
                if blue_text:
                    current_group.append(blue_text)
                else:
                    if current_group:
                        merged_text = '\n'.join(current_group)
                        self.extracted_text.append(merged_text)
                        current_group = []
            
            if current_group:
                merged_text = '\n'.join(current_group)
                self.extracted_text.append(merged_text)
            
            return self.extracted_text
        
        except Exception as e:
            print(f"âŒ è®€å–æ–‡ä»¶æ™‚ç™¼ç”ŸéŒ¯èª¤: {e}")
            sys.exit(1)
    
    def format_for_ppt(self, title="ç°¡å ±æ¨™é¡Œ"):
        """
        å°‡æå–çš„æ–‡å­—æ ¼å¼åŒ–ç‚º text_to_ppt.py å¯ç”¨çš„æ ¼å¼
        
        Args:
            title: ä¸»æ¨™é¡Œ
        
        Returns:
            str: æ ¼å¼åŒ–å¾Œçš„æ–‡å­—
        """
        if not self.extracted_text:
            return ""
        
        # åŸºæœ¬æ ¼å¼ï¼š
        # ## ä¸»æ¨™é¡Œï¼ˆè—è‰²èƒŒæ™¯ï¼‰
        # # å°æ¨™é¡Œï¼ˆç°è‰²èƒŒæ™¯ï¼‰
        # å…§å®¹è¡Œ
        
        formatted = f"## {title}\n\n"
        
        for i, text in enumerate(self.extracted_text, 1):
            # åˆ¤æ–·æ˜¯å¦ç‚ºæ¨™é¡Œï¼ˆå¯ä»¥æ ¹æ“šå¯¦éš›æƒ…æ³èª¿æ•´ï¼‰
            if len(text) < 30:  # çŸ­æ–‡å­—ç•¶ä½œå°æ¨™é¡Œ
                formatted += f"# {text}\n"
            else:  # é•·æ–‡å­—ç•¶ä½œå…§å®¹
                formatted += f"{text}\n"
        
        return formatted
    
    def save_to_file(self, output_path, title="ç°¡å ±æ¨™é¡Œ"):
        """
        å„²å­˜æå–çš„æ–‡å­—åˆ°æª”æ¡ˆï¼ˆåŒ…å«è®Šæ•¸æ¨¡æ¿ï¼‰
        
        Args:
            output_path: è¼¸å‡ºæª”æ¡ˆè·¯å¾‘
            title: ä¸»æ¨™é¡Œ
        """
        if not self.extracted_text:
            print("âš ï¸  æ²’æœ‰æ‰¾åˆ°è—è‰²æ–‡å­—")
            return False
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # å¯«å…¥è‡ªå‹•æå–çš„è®Šæ•¸
                f.write("[è®Šæ•¸]\n")
                f.write(f"æ—¥æœŸ={self.variables.get('æ—¥æœŸ', '2026å¹´1æœˆ1æ—¥')}\n")
                f.write(f"ç¦®æ‹œé¡å‹={self.variables.get('ç¦®æ‹œé¡å‹', 'é€±ä¸‰ç¦®æ‹œ')}\n")
                f.write(f"ä¸»é¡Œ={self.variables.get('ä¸»é¡Œ', 'æˆ‘æ˜¯ä¸»é¡Œ')}\n")
                f.write(f"ç¶“æ–‡ç« ç¯€={self.variables.get('ç¶“æ–‡ç« ç¯€', 'ã€ç®´è¨€27ç« 12ç¯€ã€è©©ç¯‡46ç¯‡1ç¯€ã€‘')}\n")
                
                verse_count = sum(1 for k in self.variables.keys() if k.startswith('ç¶“æ–‡'))
                if verse_count > 0:
                    for i in range(1, verse_count + 1):
                        f.write(f"ç¶“æ–‡{i}={self.variables.get(f'ç¶“æ–‡{i}', '')}\n")
                else:
                    f.write("ç¶“æ–‡1=ã€ˆç®´è¨€27ç« 12ç¯€ã€‰XXXXXXXXã€‚\n")
                    f.write("ç¶“æ–‡2=ã€ˆè©©ç¯‡46ç¯‡1ç¯€ã€‰OOOOOOOOã€‚\n")
                f.write("[è®Šæ•¸çµæŸ]\n\n")
                
                # å¯«å…¥æå–çš„è—è‰²æ–‡å­—å…§å®¹
                for text in self.extracted_text:
                    f.write(f"{text}\n\n")
            
            print(f"âœ… æˆåŠŸæå– {len(self.extracted_text)} æ®µè—è‰²æ–‡å­—")
            print(f"ğŸ“ å·²å„²å­˜åˆ°ï¼š{output_path}")
            return True
        
        except Exception as e:
            print(f"âŒ å„²å­˜æª”æ¡ˆæ™‚ç™¼ç”ŸéŒ¯èª¤: {e}")
            return False


def main():
    """ä¸»ç¨‹å¼"""
    # åƒæ•¸ 1ï¼šè¼¸å…¥ Word æª”æ¡ˆï¼ˆå¯é¸ï¼Œé è¨­ input.docxï¼‰
    input_file = sys.argv[1] if len(sys.argv) >= 2 else "input.docx"
    
    # å›ºå®šè¼¸å‡ºæª”æ¡ˆç‚º output.txt
    output_file = "output.txt"
    
    # å¾ config.txt è®€å–é¡è‰²è¨­å®šï¼ˆå¯é¸ï¼Œé è¨­è—è‰²ï¼‰
    target_color = None
    config_file = "config.txt"
    
    if os.path.exists(config_file):
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                in_color_section = False
                for line in f:
                    line = line.strip()
                    
                    if line == '[é¡è‰²è¨­å®š]':
                        in_color_section = True
                        continue
                    
                    if line.startswith('[') and line.endswith(']'):
                        in_color_section = False
                        continue
                    
                    if in_color_section and line.startswith('æå–æ–‡å­—é¡è‰²'):
                        if '=' in line:
                            _, value = line.split('=', 1)
                            value = value.strip()
                            
                            if value.startswith('#'):
                                target_color = value
                            else:
                                rgb = tuple(int(c.strip()) for c in value.split(','))
                                if len(rgb) == 3:
                                    target_color = rgb
                            break
        except Exception as e:
            print(f"âš ï¸  è­¦å‘Šï¼šè®€å– config.txt æ™‚ç™¼ç”ŸéŒ¯èª¤: {e}")
            print(f"    ä½¿ç”¨é è¨­è—è‰²")
    
    # é¡¯ç¤ºä½¿ç”¨èªªæ˜ï¼ˆå¦‚æœä½¿ç”¨ -h æˆ– --help åƒæ•¸ï¼‰
    if len(sys.argv) >= 2 and sys.argv[1] in ['-h', '--help', 'help']:
        print("ğŸ“– ç‰¹å®šé¡è‰²æ–‡å­—æå–å·¥å…·")
        print("=" * 70)
        print()
        print("ä½¿ç”¨æ–¹å¼ï¼š")
        print("  python 1_extract.py [Wordæª”æ¡ˆ]")
        print()
        print("åƒæ•¸èªªæ˜ï¼š")
        print("  Wordæª”æ¡ˆ  - Word æ–‡ä»¶è·¯å¾‘ï¼ˆé è¨­ï¼šinput.docxï¼‰")
        print()
        print("å›ºå®šè¨­å®šï¼š")
        print("  è¼¸å‡ºæª”æ¡ˆï¼šoutput.txtï¼ˆå›ºå®šï¼‰")
        print("  é¡è‰²è¨­å®šï¼šå¾ config.txt è®€å–ã€Œæå–æ–‡å­—é¡è‰²ã€ï¼ˆé è¨­ï¼šè—è‰²ï¼‰")
        print()
        print("Config é¡è‰²è¨­å®šç¯„ä¾‹ï¼ˆåœ¨ config.txt ä¸­ï¼‰ï¼š")
        print("  [é¡è‰²è¨­å®š]")
        print("  æå–æ–‡å­—é¡è‰² = 0,0,255        # è—è‰²ï¼ˆé è¨­ï¼‰")
        print("  æå–æ–‡å­—é¡è‰² = 255,0,0        # ç´…è‰²")
        print("  æå–æ–‡å­—é¡è‰² = #FF0000        # ç´…è‰²ï¼ˆ16é€²ä½ï¼‰")
        print()
        print("ç¯„ä¾‹ï¼š")
        print("  python 1_extract.py")
        print("    â†’ å¾ input.docx æå–æ–‡å­—ï¼Œè¼¸å‡ºåˆ° output.txt")
        print()
        print("  python 1_extract.py 20251231.docx")
        print("    â†’ å¾ 20251231.docx æå–æ–‡å­—ï¼Œè¼¸å‡ºåˆ° output.txt")
        print()
        print("=" * 70)
        print()
        print("ğŸ’¡ æå–å®Œæˆå¾Œï¼Œå¯ç›´æ¥åŸ·è¡Œï¼š")
        print("   python 2_generate.py")
        print()
        sys.exit(0)
    
    # æª¢æŸ¥è¼¸å…¥æª”æ¡ˆæ˜¯å¦å­˜åœ¨
    if not os.path.exists(input_file):
        print(f"âŒ éŒ¯èª¤ï¼šæ‰¾ä¸åˆ°æª”æ¡ˆ '{input_file}'")
        sys.exit(1)
    
    # åŸ·è¡Œæå–
    print(f"ğŸ“– è®€å– Word æª”æ¡ˆï¼š{input_file}")
    if target_color:
        if isinstance(target_color, str):
            print(f"ğŸ¨ ç›®æ¨™é¡è‰²ï¼š{target_color}")
        else:
            print(f"ğŸ¨ ç›®æ¨™é¡è‰²ï¼šRGB{target_color}")
    else:
        print(f"ğŸ¨ ç›®æ¨™é¡è‰²ï¼šè—è‰²ï¼ˆé è¨­ï¼‰")
    
    extractor = BlueTextExtractor(target_color=target_color, tolerance=50)
    extractor.extract_from_docx(input_file)
    
    # é¡¯ç¤ºæå–çµæœ
    if extractor.extracted_text:
        print(f"\næ‰¾åˆ° {len(extractor.extracted_text)} æ®µè—è‰²æ–‡å­—ï¼š")
        print("-" * 50)
        for i, text in enumerate(extractor.extracted_text, 1):
            preview = text[:60] + "..." if len(text) > 60 else text
            print(f"{i}. {preview}")
        print("-" * 50)
    
    # å„²å­˜çµæœ
    if extractor.save_to_file(output_file):
        print(f"\nğŸ‰ å®Œæˆï¼ç¾åœ¨å¯ä»¥åŸ·è¡Œï¼š")
        print(f"   2_generate.exe (æˆ– python 2_generate.py)")
        print(f"\næç¤ºï¼š")
        print(f"  1. è«‹å…ˆç·¨è¼¯ output.txt å¡«å…¥è®Šæ•¸")
        print(f"  2. ç„¶å¾ŒåŸ·è¡Œ 2_generate.exe ç”Ÿæˆ PPT")


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        # è¨˜éŒ„éŒ¯èª¤åˆ°æª”æ¡ˆ
        try:
            with open('error.log', 'a', encoding='utf-8') as f:
                f.write(f"\n{'='*60}\n")
                f.write(f"éŒ¯èª¤æ™‚é–“: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"ç¨‹å¼: 1_extract.py\n")
                f.write(f"éŒ¯èª¤è¨Šæ¯: {str(e)}\n")
                f.write(f"è©³ç´°è³‡è¨Š:\n{traceback.format_exc()}\n")
        except:
            pass
        
        print(f"\n{'='*60}")
        print(f"âŒ ç™¼ç”ŸéŒ¯èª¤")
        print(f"{'='*60}")
        print(f"éŒ¯èª¤è¨Šæ¯: {e}")
        print(f"\néŒ¯èª¤è©³ç´°è³‡è¨Šå·²è¨˜éŒ„åˆ° error.log")
        print(f"è«‹å°‡ error.log æä¾›çµ¦é–‹ç™¼è€…å”åŠ©é™¤éŒ¯")
        print(f"{'='*60}")
    finally:
        input("\næŒ‰ Enter éµé€€å‡º...")
